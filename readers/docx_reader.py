"""
DOCX Reader for ThreadBear

Handles .docx files using python-docx with paragraph-based segmentation.
"""
from pathlib import Path
from .registry import reader_registry


class DocxReader:
    """Reader for Microsoft Word documents."""

    @staticmethod
    def extract_text(path):
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("DOCX reading requires python-docx: pip install python-docx")
        
        doc = Document(str(path))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        return '\n\n'.join(paragraphs)

    @staticmethod
    def extract_segments(path):
        """Segment by paragraph groups."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("DOCX reading requires python-docx: pip install python-docx")
        
        doc = Document(str(path))
        segments = []
        current_segment = []
        current_tokens = 0
        max_tokens = 500
        segment_start = 0
        
        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue
            
            para_tokens = len(para.text) // 4
            
            # Start new segment if we exceed token limit
            if current_tokens + para_tokens > max_tokens and current_segment:
                segment_text = '\n\n'.join(current_segment)
                segments.append({
                    'text': segment_text,
                    'start': segment_start,
                    'end': i,
                    'tokens': len(segment_text) // 4,
                    'label': f'Section {len(segments) + 1}'
                })
                current_segment = [para.text]
                current_tokens = para_tokens
                segment_start = i
            else:
                current_segment.append(para.text)
                current_tokens += para_tokens
        
        # Save last segment
        if current_segment:
            segment_text = '\n\n'.join(current_segment)
            segments.append({
                'text': segment_text,
                'start': segment_start,
                'end': len(doc.paragraphs),
                'tokens': len(segment_text) // 4,
                'label': f'Section {len(segments) + 1}'
            })
        
        return segments


reader_registry.register(['.docx'], DocxReader, requires=['docx'])
